from fastapi import FastAPI, UploadFile, File
from models import Document
from database import SessionLocal
from PyPDF2 import PdfReader
import docx

app = FastAPI()

@app.post("/upload/")
async def upload(file: UploadFile = File(...)):
    ext = file.filename.split(".")[-1]
    content = ""

    if ext == "pdf":
        reader = PdfReader(file.file)
        content = "\n".join([page.extract_text() for page in reader.pages])
    elif ext == "docx":
        d = docx.Document(file.file)
        content = "\n".join([p.text for p in d.paragraphs])
    elif ext == "txt":
        content = (await file.read()).decode('utf-8')
    else:
        return {"error": "Unsupported file type"}

    db = SessionLocal()
    doc = Document(name=file.filename, content=content)
    db.add(doc)
    db.commit()
    db.close()

    return {"message": "File uploaded and saved."}
